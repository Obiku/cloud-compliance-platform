"""Renders PROJECT_PLAN.md to PROJECT_PLAN.docx for offline reading.

Handles the markdown subset PROJECT_PLAN.md actually uses: #/##/### headings,
'---' horizontal rules, '-' bullet lists, '|' tables, and **bold** inline runs.
Not a general-purpose markdown renderer - only covers what this one file uses.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "PROJECT_PLAN.md"
DEST = ROOT / "PROJECT_PLAN.docx"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_bold_runs(paragraph, text: str) -> None:
    pos = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render(lines: list[str], doc: Document) -> None:
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            i -= 1
            render_table(table_lines, doc)
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_bold_runs(p, line.strip()[2:])
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            add_bold_runs(p, line)
        i += 1


def render_table(table_lines: list[str], doc: Document) -> None:
    rows = [
        [cell.strip() for cell in row.strip("|").split("|")]
        for row in table_lines
        if not re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?$", row)
    ]
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            add_bold_runs(cell.paragraphs[0], cell_text)
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = SRC.read_text(encoding="utf-8").splitlines(keepends=True)
    render(lines, doc)

    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(54)

    doc.save(DEST)
    print(f"Wrote {DEST}")


if __name__ == "__main__":
    main()
